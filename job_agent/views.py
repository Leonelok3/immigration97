# job_agent/views.py
from __future__ import annotations

from datetime import timedelta
from io import BytesIO
from urllib.parse import parse_qsl, urlparse, urlunparse

from django.contrib import messages
from django.contrib.auth.decorators import login_required, user_passes_test
from django.core.exceptions import ValidationError
from django.core.validators import validate_email
from django.db import IntegrityError, transaction
from django.db.models import Q
from django.http import Http404, HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse  # ✅ AJOUT (safe)
from django.utils import timezone
from django.utils.http import urlencode
from django.utils.text import slugify
from django.views.decorators.http import require_POST
from billing.services import has_candidate_access
from outreach.ai_employer_scraper import is_direct_application_url
from outreach.models import ScrapedEmployerLead
from profiles.models import Category

from .forms import (
    CandidateDocumentsForm,
    CandidateProfileForm,
    JobLeadAddForm,
    JobLeadBulkAddForm,
    JobSearchForm,
)
from .models import (
    AnswerTemplate,
    ApplicationPack,
    CandidateDocuments,
    CandidateProfile,
    FollowUpTemplate,
    JobLead,
    JobSearch,
    LetterTemplate,
    PublicJobOffer,
)
from .services import (
    generate_application_texts,
    heuristic_match,
    render_text_template,
    send_followup_email,
)


# ======================================================
# Helpers (internal)
# ======================================================
SECTOR_CATEGORY = {
    "agriculture": "Agriculture & Agroalimentaire",
    "construction": "BTP & Construction",
    "tech": "Développement Web & IT",
    "sante": "Aide-soignant & Santé",
    "logistique": "Transport & Logistique",
    "hotellerie": "Hôtellerie & Service client",
    "education": "Éducation & Formation",
    "finance": "Comptabilité & Finance",
    "industrie": "Industrie & Manufacture",
    "commerce": "Commerce & Vente",
    "services": "Services aux entreprises",
    "autre": "Autres métiers",
}


def _staff_required(user):
    return bool(user and user.is_authenticated and user.is_staff)


def _category_for_sector(sector: str):
    name = SECTOR_CATEGORY.get((sector or "autre").strip().lower(), SECTOR_CATEGORY["autre"])
    existing = Category.objects.filter(name=name).first()
    if existing:
        return existing
    return Category.objects.get_or_create(slug=slugify(name)[:100], defaults={"name": name})[0]


def _public_offer_defaults_from_scraped_lead(lead: ScrapedEmployerLead, *, manual: bool = False) -> dict:
    details = (lead.raw_data or {}).get("job_details") or {}
    description = "\n\n".join(
        part
        for part in [
            (lead.evidence_text or "").strip(),
            (lead.verification_notes or "").strip(),
            f"Signal visa/recrutement international: {lead.visa_signal}".strip()
            if lead.visa_signal
            else "",
        ]
        if part
    )
    score = int(details.get("foreign_access_score") or lead.confidence_score or 0)
    label = (details.get("foreign_access_label") or "").strip()
    if manual and score < 70:
        label = label or "Validée manuellement"
        score = max(score, 50)

    return {
        "source": "Validation admin Immigration97" if manual else "Agent web Immigration97",
        "title": (lead.title or "Offre")[:220],
        "company": (lead.company_name or "Employeur à vérifier")[:220],
        "location": (lead.location or lead.country or "Canada")[:220],
        "category": _category_for_sector(lead.sector),
        "skills_keywords": ", ".join(lead.verification_signals or [])[:300],
        "foreign_access_score": score,
        "foreign_access_label": (label or "À vérifier")[:80],
        "salary": (details.get("salary") or "")[:180],
        "province": (details.get("province") or "")[:80],
        "application_deadline": (details.get("application_deadline") or "")[:80],
        "vacancies": (details.get("vacancies") or "")[:80],
        "who_can_apply": details.get("who_can_apply") or "",
        "description_text": description or lead.title or "Offre détectée par l'agent Immigration97.",
        "is_active": True,
    }


def _latest_search(user):
    return JobSearch.objects.filter(user=user).order_by("-created_at").first()


def _safe_get_line(block: str, prefix: str) -> str:
    """
    Cherche une ligne qui commence par 'prefix:' (case-insensitive) et retourne la valeur.
    Exemple: prefix='URL' récupère 'URL: https://...'
    """
    if not block:
        return ""
    for line in block.splitlines():
        line = line.strip()
        if line.lower().startswith(prefix.lower() + ":"):
            return line.split(":", 1)[1].strip()
    return ""


def _extract_description(block: str) -> str:
    """
    Récupère tout ce qui suit 'Description:' dans un bloc.
    """
    if not block:
        return ""
    low = block.lower()
    if "description:" not in low:
        return ""
    idx = low.index("description:")
    return block[idx + len("description:") :].strip()


def _render_letter_from_template(
    template_text: str, *, title: str, company: str, location: str, name: str
) -> str:
    """
    Remplit {title} {company} {location} {name} si présents.
    Si un placeholder est invalide, renvoie le texte brut.
    """
    data = {
        "title": title or "",
        "company": company or "",
        "location": location or "",
        "name": name or "",
    }
    try:
        return (template_text or "").format(**data).strip()
    except Exception:
        return (template_text or "").strip()


def _answers_from_admin_templates(language: str) -> dict:
    """
    Récupère AnswerTemplate actifs pour la langue demandée.
    Retour: { key: content }
    """
    lang = (language or "fr").lower()
    qs = AnswerTemplate.objects.filter(is_active=True, language=lang).order_by("key", "id")
    out: dict[str, str] = {}
    for t in qs:
        out.setdefault(t.key, t.content)
    return out


def _letter_template_for_language(language: str):
    lang = (language or "fr").lower()
    return LetterTemplate.objects.filter(is_active=True, language=lang).order_by("-id").first()


def _followup_template_for_language(language: str):
    lang = (language or "fr").lower()
    return FollowUpTemplate.objects.filter(is_active=True, language=lang).order_by("-id").first()

def build_or_update_pack(*, user, lead, profile, docs) -> ApplicationPack:
    """
    Construit ou met à jour le pack candidature pour une offre (lead).
    SAFE: reprend exactement la logique existante du bouton "generate_pack".
    """
    language = (lead.search.language if lead.search else (profile.language or "fr")) or "fr"
    language = (language or "fr").lower()

    offer_title = lead.title or (lead.search.title if lead.search else "Poste")
    company = lead.company or ""
    location = lead.location or ""
    name = profile.full_name or user.get_username()

    from .application_coach import generate_tailored_application

    result = generate_tailored_application(
        offer_title=offer_title,
        company=company,
        location=location,
        offer_text=lead.description_text or "",
        cv_text=docs.cv_text or "",
        base_letter=docs.base_letter_text or "",
        language=language,
        candidate_name=name,
    )

    answers = result.suggested_answers or {}
    admin_answers = _answers_from_admin_templates(language)
    if admin_answers:
        merged = dict(answers)
        for k, v in admin_answers.items():
            merged[k] = v
        answers = merged

    # Enregistrer / mettre à jour le pack
    pack, _ = ApplicationPack.objects.get_or_create(user=user, lead=lead)
    pack.generated_letter = (result.generated_letter or "").strip()
    pack.suggested_answers = answers or {}
    pack.email_subject = (result.email_subject or "").strip()
    pack.generated_email = (result.generated_email or "").strip()
    pack.tailored_cv_text = (result.tailored_cv_text or "").strip()
    pack.ats_score = result.ats_score or 0
    pack.matched_keywords = result.matched_keywords or []
    pack.missing_keywords = result.missing_keywords or []
    pack.coach_notes = (result.coach_notes or "").strip()
    pack.ai_status = result.ai_status or ""
    pack.ai_generated_at = timezone.now()
    pack.save()

    return pack

def _menu_pack_lead_id_for_user(user) -> int | None:
    """
    ✅ Safe helper: permet au menu de générer un lien pack (qui exige lead_id)
    sans casser si l'utilisateur n'a aucune offre.
    """
    lead = JobLead.objects.filter(user=user).order_by("-updated_at", "-created_at").only("id").first()
    return lead.id if lead else None


# ======================================================
# Pages principales
# ======================================================
@login_required
def dashboard(request):
    profile, _ = CandidateProfile.objects.get_or_create(user=request.user)
    docs, _ = CandidateDocuments.objects.get_or_create(user=request.user)

    searches = JobSearch.objects.filter(user=request.user).order_by("-created_at")[:10]
    leads = JobLead.objects.filter(user=request.user).order_by("-created_at")[:20]

    counts = {
        "found": JobLead.objects.filter(user=request.user, status=JobLead.STATUS_FOUND).count(),
        "to_apply": JobLead.objects.filter(user=request.user, status=JobLead.STATUS_TO_APPLY).count(),
        "applied": JobLead.objects.filter(user=request.user, status=JobLead.STATUS_APPLIED).count(),
        "followup": JobLead.objects.filter(user=request.user, status=JobLead.STATUS_FOLLOWUP).count(),
        "reply": JobLead.objects.filter(user=request.user, status=JobLead.STATUS_REPLY).count(),
    }

    # ✅ AJOUTS SAFE POUR LE FRONT (ne casse rien)
    menu_pack_lead_id = _menu_pack_lead_id_for_user(request.user)

    return render(
        request,
        "job_agent/dashboard.html",
        {
            "profile": profile,
            "docs": docs,
            "searches": searches,
            "leads": leads,
            "counts": counts,
            # ✅
            "menu_pack_lead_id": menu_pack_lead_id,
            "dashboard_url": reverse("job_agent:dashboard"),
            "leads_url": reverse("job_agent:lead_list"),
            "kanban_url": reverse("job_agent:kanban"),
            "lead_add_url": reverse("job_agent:lead_add"),
        },
    )


@login_required
def profile_edit(request):
    profile, _ = CandidateProfile.objects.get_or_create(user=request.user)

    if request.method == "POST":
        form = CandidateProfileForm(request.POST, instance=profile)
        if form.is_valid():
            form.save()
            messages.success(request, "Profil candidat mis à jour.")
            return redirect("job_agent:dashboard")
    else:
        form = CandidateProfileForm(instance=profile)

    # ✅ menu safe
    return render(
        request,
        "job_agent/profile_form.html",
        {
            "form": form,
            "menu_pack_lead_id": _menu_pack_lead_id_for_user(request.user),
        },
    )


@login_required
def documents_edit(request):
    docs, _ = CandidateDocuments.objects.get_or_create(user=request.user)

    if request.method == "POST":
        form = CandidateDocumentsForm(request.POST, request.FILES, instance=docs)
        if form.is_valid():
            obj = form.save(commit=False)

            # ✅ Extraction automatique du texte du CV PDF
            auto_extract = form.cleaned_data.get("auto_extract_cv", True)
            uploaded_cv = request.FILES.get("cv_file")

            if auto_extract and uploaded_cv:
                from .services import extract_cv_text_from_file

                extracted = extract_cv_text_from_file(uploaded_cv, filename=uploaded_cv.name)

                if extracted.strip():
                    obj.cv_text = extracted
                    messages.success(request, "Texte CV extrait automatiquement depuis le PDF.")
                else:
                    messages.warning(
                        request,
                        "Impossible d’extraire le texte du PDF (scan image ou PDF protégé). "
                        "Dans ce cas, colle le texte du CV manuellement."
                    )

            obj.save()
            messages.success(request, "Documents mis à jour.")
            return redirect("job_agent:documents_edit")
    else:
        form = CandidateDocumentsForm(instance=docs)

    from .application_coach import analyze_cv_quality

    cv_quality = analyze_cv_quality(docs.cv_text or "")

    # ✅ menu safe
    return render(
        request,
        "job_agent/documents_form.html",
        {
            "form": form,
            "docs": docs,
            "cv_quality": cv_quality,
            "menu_pack_lead_id": _menu_pack_lead_id_for_user(request.user),
        },
    )


@login_required
def search_create(request):
    if request.method == "POST":
        form = JobSearchForm(request.POST)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.user = request.user
            obj.save()
            messages.success(request, "Recherche créée.")
            return redirect("job_agent:lead_list")
    else:
        form = JobSearchForm()

    # ✅ menu safe
    return render(
        request,
        "job_agent/search_form.html",
        {
            "form": form,
            "menu_pack_lead_id": _menu_pack_lead_id_for_user(request.user),
        },
    )


# ======================================================
# Offres (leads)
# ======================================================
@login_required
def lead_list(request):
    leads = JobLead.objects.filter(user=request.user).order_by("-created_at")
    searches = JobSearch.objects.filter(user=request.user).order_by("-created_at")

    # ✅ menu safe
    return render(
        request,
        "job_agent/lead_list.html",
        {
            "leads": leads,
            "searches": searches,
            "menu_pack_lead_id": _menu_pack_lead_id_for_user(request.user),
            "dashboard_url": reverse("job_agent:dashboard"),
            "all_leads_url": reverse("job_agent:lead_list"),
        },
    )


@login_required
def lead_add(request):
    if request.method == "POST":
        form = JobLeadAddForm(request.POST)
        if form.is_valid():
            lead = form.save(commit=False)
            lead.user = request.user

            # Si l'utilisateur ne choisit pas de recherche, on met la dernière
            if not lead.search:
                lead.search = _latest_search(request.user)

            try:
                lead.save()
            except IntegrityError:
                messages.warning(request, "Tu as déjà ajouté cette offre (même URL).")
                return redirect("job_agent:lead_list")

            # ✅ AUTO-SCORE DIRECT (gain de temps)
            docs, _ = CandidateDocuments.objects.get_or_create(user=request.user)

            if (lead.description_text or "").strip():
                keywords = lead.search.keywords if lead.search else ""
                try:
                    # IA sémantique si dispo (optionnel)
                    from .ai_matching import semantic_match  # type: ignore

                    score = int(semantic_match(docs.cv_text or "", lead.description_text or ""))
                    lead.match_score = max(0, min(score, 100))
                    lead.match_summary = "Scoring sémantique IA (embeddings)."
                except Exception:
                    res = heuristic_match(docs.cv_text or "", lead.description_text or "", keywords=keywords)
                    lead.match_score = res.score
                    lead.match_summary = res.summary

                if lead.match_score >= 60 and lead.status == JobLead.STATUS_FOUND:
                    lead.status = JobLead.STATUS_TO_APPLY

                lead.save(update_fields=["match_score", "match_summary", "status"])
                messages.success(request, f"Offre ajoutée + scorée automatiquement : {lead.match_score}/100")
            else:
                messages.success(
                    request,
                    "Offre ajoutée. Colle la description de l’offre pour obtenir un scoring automatique.",
                )

            return redirect("job_agent:lead_detail", lead_id=lead.id)
    else:
        form = JobLeadAddForm()
        last_search = _latest_search(request.user)
        if last_search:
            form.initial["search"] = last_search.id

    return render(
        request,
        "job_agent/lead_add.html",
        {
            "form": form,
            "menu_pack_lead_id": _menu_pack_lead_id_for_user(request.user),
        },
    )


from django.core import signing  # Ajout de l'import pour le token
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.db import transaction
from django.utils import timezone
from datetime import timedelta
from django.core.validators import validate_email
from django.core.exceptions import ValidationError
from django.contrib import messages

@login_required
@transaction.atomic
def lead_detail(request, lead_id: int):
    lead = get_object_or_404(JobLead, id=lead_id, user=request.user)
    docs, _ = CandidateDocuments.objects.get_or_create(user=request.user)
    profile, _ = CandidateProfile.objects.get_or_create(user=request.user)

    def compute_score() -> tuple[int, str]:
        cv_text = (docs.cv_text or "").strip()
        offer_text = (lead.description_text or "").strip()
        keywords = lead.search.keywords if lead.search else ""

        if not offer_text:
            return 0, "Aucune description d'offre. Colle la description pour un scoring précis."

        # 1) IA sémantique si dispo
        try:
            from .ai_matching import semantic_match  # type: ignore

            score = semantic_match(cv_text, offer_text)
            return int(score), "Scoring sémantique IA (embeddings)."
        except Exception:
            # 2) fallback heuristique
            res = heuristic_match(cv_text, offer_text, keywords=keywords)
            return res.score, res.summary

    if request.method == "POST":
        action = (request.POST.get("action") or "").strip()

        # ======================================================
        # ✅ 1) SAVE CONTACT EMAIL (SAFE)
        # ======================================================
        if action == "save_contact_email":
            email = (request.POST.get("contact_email") or "").strip()

            if not email:
                messages.error(request, "Ajoute un email recruteur.")
                return redirect("job_agent:lead_detail", lead_id=lead.id)

            try:
                validate_email(email)
            except ValidationError:
                messages.error(request, "Email invalide.")
                return redirect("job_agent:lead_detail", lead_id=lead.id)

            lead.contact_email = email
            lead.save(update_fields=["contact_email"])
            messages.success(request, "Email recruteur enregistré ✅")
            return redirect("job_agent:lead_detail", lead_id=lead.id)

        # ======================================================
        # ✅ 2) SEND FOLLOWUP NOW (SAFE + ANTI-SPAM + TEMPLATE CHECK)
        # ======================================================
        if action == "send_followup_now":
            email = (request.POST.get("contact_email") or lead.contact_email or "").strip()

            if not email:
                messages.error(request, "Ajoute d’abord l’email recruteur.")
                return redirect("job_agent:lead_detail", lead_id=lead.id)

            try:
                validate_email(email)
            except ValidationError:
                messages.error(request, "Email invalide.")
                return redirect("job_agent:lead_detail", lead_id=lead.id)

            # Anti-spam : 24h minimum entre 2 relances
            if lead.followup_sent_at and (timezone.now() - lead.followup_sent_at) < timedelta(hours=24):
                messages.warning(request, "Relance déjà envoyée récemment. Attends 24h.")
                return redirect("job_agent:lead_detail", lead_id=lead.id)

            language = (lead.search.language if lead.search else (profile.language or "fr")).lower()
            tpl = _followup_template_for_language(language)

            name = profile.full_name or request.user.get_username()
            title = lead.title or (lead.search.title if lead.search else "Poste")
            company = lead.company or ""
            location = lead.location or ""

            if tpl:
                subject = render_text_template(
                    tpl.subject, name=name, title=title, company=company, location=location
                )
                body = render_text_template(
                    tpl.content, name=name, title=title, company=company, location=location
                )
            else:
                subject = f"Relance — {title} ({company})"
                body = (
                    "Bonjour,\n\n"
                    f"Je me permets de relancer ma candidature au poste {title}"
                    f"{(' à ' + location) if location else ''}.\n"
                    "Je reste disponible pour un échange (entretien / test).\n\n"
                    f"Cordialement,\n{name}\n"
                )

            # Sécurité: empêcher sujet/corps vide
            subject = (subject or "").strip()
            body = (body or "").strip()
            if not subject or not body:
                messages.error(
                    request,
                    "Template de relance invalide (objet/corps vide). Corrige dans l’admin.",
                )
                return redirect("job_agent:lead_detail", lead_id=lead.id)

            # Envoi SMTP
            try:
                send_followup_email(to_email=email, subject=subject, body=body)
            except Exception as e:
                messages.error(request, f"Erreur SMTP : {e}")
                return redirect("job_agent:lead_detail", lead_id=lead.id)

            # Update lead
            lead.contact_email = email
            lead.followup_sent_at = timezone.now()
            lead.status = JobLead.STATUS_FOLLOWUP
            lead.save(update_fields=["contact_email", "followup_sent_at", "status"])

            messages.success(request, "Relance envoyée ✅")
            return redirect("job_agent:lead_detail", lead_id=lead.id)

        # ======================================================
        # ✅ 3) SCORE
        # ======================================================
        if action == "score":
            score, summary = compute_score()
            lead.match_score = score
            lead.match_summary = summary

            if lead.match_score >= 60 and lead.status == JobLead.STATUS_FOUND:
                lead.status = JobLead.STATUS_TO_APPLY

            lead.save(update_fields=["match_score", "match_summary", "status"])
            messages.success(request, f"Score mis à jour : {lead.match_score}/100")
            return redirect("job_agent:lead_detail", lead_id=lead.id)

        # ======================================================
        # ✅ 4) GENERATE PACK (EMAIL + LETTRE + REPONSES)
        # ======================================================
        if action == "generate_pack":
            build_or_update_pack(user=request.user, lead=lead, profile=profile, docs=docs)
            messages.success(request, "Pack candidature généré (email + lettre + réponses) ✅")
            return redirect("job_agent:pack_detail", lead_id=lead.id)

        # ======================================================
        # ✅ 5) SET STATUS (auto applied_at)
        # ======================================================
        if action == "set_status":
            new_status = request.POST.get("status")
            valid = {c[0] for c in JobLead.STATUS_CHOICES}

            if new_status in valid:
                lead.status = new_status

                # Auto: date candidature si "Postulée"
                if new_status == JobLead.STATUS_APPLIED and not lead.applied_at:
                    lead.applied_at = timezone.now()

                lead.save(update_fields=["status", "applied_at"])
                messages.success(request, "Statut mis à jour ✅")
            else:
                messages.error(request, "Statut invalide.")

            return redirect("job_agent:lead_detail", lead_id=lead.id)

    # --- GÉNÉRATION DU TOKEN POUR AUTOFILL ---
    imm97_token = signing.dumps({"uid": request.user.id, "lead_id": lead.id},salt="imm97_autofill")


    # ✅ IMPORTANT: on renvoie aussi profile + status_choices + menu_pack_lead_id + imm97_token
    return render(
        request,
        "job_agent/lead_detail.html",
        {
            "lead": lead,
            "docs": docs,
            "profile": profile,
            "status_choices": JobLead.STATUS_CHOICES,
            "menu_pack_lead_id": lead.id,
            "imm97_token": imm97_token,  # Ajout au contexte
        },
    )
def _premium_agent_required(request):
    messages.error(
        request,
        "L'Agent IA Candidature est réservé aux comptes Premium Candidat: CV adapté, lettre, email, score ATS et instructions de candidature.",
    )
    target = f"{reverse('billing:pricing')}?{urlencode({'next': request.get_full_path()})}"
    return redirect(target)


@login_required
def pack_detail(request, lead_id: int):
    lead = get_object_or_404(JobLead, id=lead_id, user=request.user)
    if not has_candidate_access(request.user):
        return _premium_agent_required(request)
    pack = get_object_or_404(ApplicationPack, lead=lead, user=request.user)

    return render(
        request,
        "job_agent/pack_detail.html",
        {
            "lead": lead,
            "pack": pack,
            "menu_pack_lead_id": lead.id,
            "leads_url": reverse("job_agent:lead_list"),
            "lead_detail_url": reverse("job_agent:lead_detail", kwargs={"lead_id": lead.id}),
        },
    )


# ======================================================
# Offres publiques (admin -> users)
# ======================================================
def _infer_public_offer_filters(ai_query: str, categories):
    text = (ai_query or "").strip().lower()
    inferred = {
        "country": "",
        "category_id": "",
        "query": "",
        "notes": [],
    }
    if not text:
        return inferred

    country_aliases = {
        "canada": "Canada",
        "québec": "Canada",
        "quebec": "Canada",
        "montréal": "Canada",
        "montreal": "Canada",
        "france": "France",
        "paris": "France",
        "allemagne": "Allemagne",
        "germany": "Allemagne",
        "deutschland": "Allemagne",
        "ausbildung": "Allemagne",
        "espagne": "Espagne",
        "spain": "Espagne",
        "espana": "Espagne",
        "españa": "Espagne",
        "belgique": "Belgique",
        "australie": "Australie",
        "australia": "Australie",
        "nouvelle-zélande": "Nouvelle-Zélande",
        "nouvelle zelande": "Nouvelle-Zélande",
        "new zealand": "Nouvelle-Zélande",
        "europe": "Europe",
        "remote": "Remote",
        "télétravail": "Remote",
        "teletravail": "Remote",
    }
    for token, country in country_aliases.items():
        if token in text:
            inferred["country"] = country
            inferred["notes"].append(f"Pays détecté: {country}")
            break

    category_aliases = {
        "santé": ["santé", "aide-soignant", "aide soignant", "infirmier", "caregiver", "nurse"],
        "btp": ["btp", "construction", "maçon", "macon", "coffreur", "chantier", "bâtiment", "batiment"],
        "it": ["développeur", "developpeur", "python", "django", "web", "wordpress", "support web", "it", "informatique"],
        "logistique": ["chauffeur", "driver", "logistique", "camion", "livreur", "transport"],
        "hôtellerie": ["hotel", "hôtel", "hotellerie", "hôtellerie", "restaurant", "serveur", "cuisine", "service client"],
        "finance": ["comptable", "finance", "excel", "assistant comptable", "paie", "accounting"],
        "agriculture": ["agriculture", "ferme", "ouvrier agricole", "agricole", "farm"],
        "commerce": ["vente", "commercial", "vendeur", "sales", "commerce"],
        "éducation": ["enseignant", "education", "éducation", "formation", "teacher"],
        "ausbildung": ["ausbildung", "duale", "berufsausbildung", "apprentissage", "formation professionnelle"],
    }
    category_list = list(categories)
    for _group, aliases in category_aliases.items():
        if any(alias in text for alias in aliases):
            for category in category_list:
                name = category.name.lower()
                if any(alias in name for alias in aliases) or _group in name:
                    inferred["category_id"] = str(category.id)
                    inferred["notes"].append(f"Métier détecté: {category.name}")
                    break
            if inferred["category_id"]:
                break

    cleaned = text
    for word in [
        "trouve-moi", "trouve moi", "cherche", "recherche", "offres", "offre",
        "emploi", "emplois", "poste", "postes", "avec visa", "visa sponsor",
        "sponsorship", "pour africain", "africain", "africaine",
    ]:
        cleaned = cleaned.replace(word, " ")
    cleaned = " ".join(part for part in cleaned.split() if len(part) > 2)
    inferred["query"] = cleaned[:120]
    if "visa" in text or "sponsor" in text or "lmia" in text or "ausbildung" in text:
        inferred["notes"].append("Priorité: offres avec signaux visa/sponsorship si disponibles")
    return inferred


def _public_offer_opportunity_type(offer: PublicJobOffer) -> dict:
    text = " ".join(
        part
        for part in [
            offer.title,
            offer.company,
            offer.location,
            offer.foreign_access_label,
            offer.skills_keywords,
            offer.description_text,
            offer.url,
        ]
        if part
    ).lower()
    training_signals = [
        "ausbildung",
        "duales studium",
        "berufsausbildung",
        "apprentissage",
        "contrat de formation",
        "formation professionnelle",
        "apprenticeship",
        "azubi",
        "lehrstelle",
        "aubi-plus",
        "azubiyo",
    ]
    if any(signal in text for signal in training_signals):
        return {
            "code": "formation",
            "label": "Contrat de formation",
            "badge": "Formation payée",
            "hint": "Ausbildung / apprentissage: le candidat apprend un metier et recoit souvent une remuneration mensuelle.",
        }
    return {
        "code": "emploi",
        "label": "Emploi direct",
        "badge": "Emploi",
        "hint": "Poste a pourvoir: le candidat postule comme travailleur sur une offre d'emploi classique.",
    }


def public_offers(request):
    offers = (
        PublicJobOffer.objects.filter(is_active=True)
        .exclude(url="")
        .select_related("category")
        .order_by("-foreign_access_score", "-updated_at", "-created_at")
    )
    selected_country = (request.GET.get("country") or "").strip()
    selected_category = (request.GET.get("category") or "").strip()
    selected_type = (request.GET.get("type") or "").strip()
    query = (request.GET.get("q") or "").strip()
    ai_query = (request.GET.get("ai_query") or "").strip()
    categories = Category.objects.all().order_by("name")

    inferred = _infer_public_offer_filters(ai_query, categories)
    if ai_query:
        selected_country = selected_country or inferred["country"]
        selected_category = selected_category or inferred["category_id"]
        query = query or inferred["query"]

    if selected_country:
        if selected_country == "Allemagne":
            offers = offers.filter(
                Q(location__icontains=selected_country)
                | Q(url__icontains="arbeitsagentur.de")
                | Q(url__icontains="make-it-in-germany.com")
                | Q(url__icontains="ausbildung.de")
                | Q(url__icontains="azubiyo.de")
                | Q(url__icontains="aubi-plus.de")
                | Q(description_text__icontains="Ausbildung")
            )
        elif selected_country == "Canada":
            offers = offers.filter(
                Q(location__icontains=selected_country)
                | Q(location__icontains="CA")
                | Q(url__icontains="jobbank.gc.ca")
                | Q(url__icontains="guichetemplois.gc.ca")
            )
        elif selected_country == "Espagne":
            offers = offers.filter(
                Q(location__icontains=selected_country)
                | Q(location__icontains="ES")
                | Q(location__icontains="Spain")
                | Q(location__icontains="España")
                | Q(url__icontains="eures.europa.eu")
                | Q(url__icontains="sepe.es")
                | Q(url__icontains="empleate.gob.es")
                | Q(url__icontains="sistemanacionalempleo.es")
                | Q(url__icontains="infojobs.net")
                | Q(url__icontains="turijobs.com")
                | Q(url__icontains="hosteleo.com")
                | Q(description_text__icontains="agricultura")
                | Q(description_text__icontains="hosteleria")
                | Q(description_text__icontains="construccion")
                | Q(description_text__icontains="contratacion en origen")
            )
        else:
            offers = offers.filter(location__icontains=selected_country)
    if selected_category:
        offers = offers.filter(category_id=selected_category)
    if query:
        offers = offers.filter(
            Q(title__icontains=query)
            | Q(company__icontains=query)
            | Q(location__icontains=query)
            | Q(skills_keywords__icontains=query)
            | Q(description_text__icontains=query)
        )

    if selected_type == "formation":
        offers = offers.filter(
            Q(title__icontains="Ausbildung")
            | Q(title__icontains="Duales Studium")
            | Q(title__icontains="Berufsausbildung")
            | Q(title__icontains="Apprentissage")
            | Q(foreign_access_label__icontains="Ausbildung")
            | Q(skills_keywords__icontains="ausbildung")
            | Q(description_text__icontains="Ausbildung")
            | Q(description_text__icontains="Duales Studium")
            | Q(description_text__icontains="Berufsausbildung")
            | Q(url__icontains="ausbildung.de")
            | Q(url__icontains="azubiyo.de")
            | Q(url__icontains="aubi-plus.de")
            | Q(url__icontains="angebotsart=4")
        )
    elif selected_type == "emploi":
        offers = offers.exclude(
            Q(title__icontains="Ausbildung")
            | Q(title__icontains="Duales Studium")
            | Q(title__icontains="Berufsausbildung")
            | Q(title__icontains="Apprentissage")
            | Q(foreign_access_label__icontains="Ausbildung")
            | Q(skills_keywords__icontains="ausbildung")
            | Q(description_text__icontains="Ausbildung")
            | Q(description_text__icontains="Duales Studium")
            | Q(description_text__icontains="Berufsausbildung")
            | Q(url__icontains="ausbildung.de")
            | Q(url__icontains="azubiyo.de")
            | Q(url__icontains="aubi-plus.de")
            | Q(url__icontains="angebotsart=4")
        )

    countries = [
        ("Canada", "Canada"),
        ("Remote", "Remote"),
        ("France", "France"),
        ("Allemagne", "Allemagne"),
        ("Espagne", "Espagne"),
        ("Belgique", "Belgique"),
        ("Australie", "Australie"),
        ("Nouvelle-Zélande", "Nouvelle-Zélande"),
        ("Europe", "Europe"),
    ]
    foreign_candidate_signals = (
        Q(foreign_access_score__gte=55)
        | Q(foreign_access_label__icontains="étranger")
        | Q(foreign_access_label__icontains="foreign")
        | Q(foreign_access_label__icontains="visa")
        | Q(foreign_access_label__icontains="LMIA")
        | Q(foreign_access_label__icontains="EIMT")
        | Q(who_can_apply__icontains="foreign")
        | Q(who_can_apply__icontains="international")
        | Q(who_can_apply__icontains="étranger")
        | Q(description_text__icontains="foreign workers")
        | Q(description_text__icontains="international applicants")
        | Q(description_text__icontains="visa sponsorship")
        | Q(description_text__icontains="LMIA")
        | Q(description_text__icontains="EIMT")
        | Q(description_text__icontains="Afrique")
        | Q(description_text__icontains="African")
        | Q(description_text__icontains="Cameroon")
        | Q(description_text__icontains="Cameroun")
        | Q(skills_keywords__icontains="visa")
        | Q(skills_keywords__icontains="foreign")
        | Q(skills_keywords__icontains="international")
        | Q(url__icontains="jobbank.gc.ca")
        | Q(url__icontains="guichetemplois.gc.ca")
        | Q(url__icontains="make-it-in-germany.com")
        | Q(url__icontains="arbeitsagentur.de")
        | Q(url__icontains="eures.europa.eu")
    )
    offers = offers.filter(foreign_candidate_signals)
    offers = _dedupe_public_offers(offers)
    offers_count = len(offers)
    for offer in offers:
        opportunity_type = _public_offer_opportunity_type(offer)
        offer.opportunity_type = opportunity_type["code"]
        offer.opportunity_type_label = opportunity_type["label"]
        offer.opportunity_type_badge = opportunity_type["badge"]
        offer.opportunity_type_hint = opportunity_type["hint"]
        raw_keywords = (offer.skills_keywords or "").replace("[", "").replace("]", "").replace("'", "").replace('"', "")
        offer.keyword_list = [
            item.strip()
            for item in raw_keywords.split(",")
            if len(item.strip()) > 2
        ][:6]
        parsed_url = urlparse(offer.url or "")
        host = parsed_url.netloc.removeprefix("www.")
        offer.website_url = f"{parsed_url.scheme}://{parsed_url.netloc}" if parsed_url.scheme and parsed_url.netloc else offer.url
        offer.website_label = host or "Site officiel"
        offer.contact_display = "Postuler sur la page officielle"
        offer.direct_apply_label = "Postuler sur l'offre officielle"
        offer.foreign_focus_note = _foreign_focus_note(offer)

    return render(
        request,
        "job_agent/public_offers.html",
        {
            "offers": offers,
            "offers_count": offers_count,
            "categories": categories,
            "countries": countries,
            "selected_country": selected_country,
            "selected_category": selected_category,
            "selected_type": selected_type,
            "opportunity_types": [
                ("emploi", "Emploi direct"),
                ("formation", "Contrat de formation"),
            ],
            "query": query,
            "ai_query": ai_query,
            "ai_notes": inferred["notes"],
            "has_candidate_premium": has_candidate_access(request.user)
            if request.user.is_authenticated
            else False,
            "menu_pack_lead_id": _menu_pack_lead_id_for_user(request.user)
            if request.user.is_authenticated
            else None,
        },
    )


def _dedupe_public_offers(offers):
    seen = set()
    unique_offers = []
    for offer in offers:
        if not is_direct_application_url(offer.url):
            continue
        key = _canonical_offer_url(offer.url)
        if not key or key in seen:
            continue
        seen.add(key)
        unique_offers.append(offer)
    return unique_offers


def _canonical_offer_url(url: str) -> str:
    parsed = urlparse(url or "")
    if not parsed.scheme or not parsed.netloc:
        return (url or "").strip().lower()
    ignored_prefixes = ("utm_",)
    ignored_keys = {"fbclid", "gclid", "msclkid", "source", "ref", "referrer"}
    query = [
        (key, value)
        for key, value in parse_qsl(parsed.query, keep_blank_values=True)
        if key.lower() not in ignored_keys and not key.lower().startswith(ignored_prefixes)
    ]
    normalized = parsed._replace(
        scheme=parsed.scheme.lower(),
        netloc=parsed.netloc.lower().removeprefix("www."),
        query=urlencode(query, doseq=True),
        fragment="",
    )
    return urlunparse(normalized).rstrip("/").lower()


def _foreign_focus_note(offer: PublicJobOffer) -> str:
    text = " ".join(
        [
            offer.foreign_access_label or "",
            offer.who_can_apply or "",
            offer.description_text or "",
            offer.skills_keywords or "",
            offer.url or "",
        ]
    ).lower()
    if any(term in text for term in ("cameroon", "cameroun")):
        return "Signal Cameroun détecté"
    if any(term in text for term in ("africa", "african", "afrique", "africain")):
        return "Signal candidats africains"
    if any(term in text for term in ("visa", "lmia", "eimt", "foreign", "international", "étranger")):
        return "Signal accès étranger"
    return "À vérifier avant candidature"


@login_required
@user_passes_test(_staff_required)
def review_detected_offers(request):
    public_urls = PublicJobOffer.objects.values_list("url", flat=True)
    leads = (
        ScrapedEmployerLead.objects.exclude(job_url__in=public_urls)
        .exclude(status="imported")
        .order_by("-last_seen_at", "-verification_score", "-confidence_score")
    )

    source = (request.GET.get("source") or "").strip()
    decision = (request.GET.get("decision") or "").strip()
    q = (request.GET.get("q") or "").strip()

    if source:
        leads = leads.filter(job_url__icontains=source)
    if decision:
        leads = leads.filter(verification_decision=decision)
    if q:
        leads = leads.filter(
            Q(title__icontains=q)
            | Q(company_name__icontains=q)
            | Q(location__icontains=q)
            | Q(evidence_text__icontains=q)
            | Q(job_url__icontains=q)
        )

    leads = list(leads[:120])
    for lead in leads:
        details = (lead.raw_data or {}).get("job_details") or {}
        lead.foreign_access_score = details.get("foreign_access_score") or lead.confidence_score or 0
        lead.foreign_access_label = details.get("foreign_access_label") or "À vérifier"
        lead.salary_display = details.get("salary") or ""
        lead.province_display = details.get("province") or ""
        lead.deadline_display = details.get("application_deadline") or ""
        lead.vacancies_display = details.get("vacancies") or ""
        lead.who_can_apply_display = details.get("who_can_apply") or ""
        parsed = urlparse(lead.job_url or "")
        lead.source_host = parsed.netloc.removeprefix("www.")
        lead.direct_link_ok = is_direct_application_url(lead.job_url)

    return render(
        request,
        "job_agent/review_detected_offers.html",
        {
            "leads": leads,
            "total_count": len(leads),
            "source": source,
            "decision": decision,
            "q": q,
        },
    )


@login_required
@user_passes_test(_staff_required)
@require_POST
@transaction.atomic
def publish_detected_offer(request, lead_id: int):
    lead = get_object_or_404(ScrapedEmployerLead, id=lead_id)
    if not is_direct_application_url(lead.job_url):
        messages.error(request, "Publication refusée: le lien n'est pas une annonce directe.")
        return redirect("job_agent:review_detected_offers")

    _offer, created = PublicJobOffer.objects.update_or_create(
        url=lead.job_url,
        defaults=_public_offer_defaults_from_scraped_lead(lead, manual=True),
    )
    lead.status = "imported"
    lead.save(update_fields=["status", "updated_at"])

    if created:
        messages.success(request, "Offre publiée manuellement dans les offres publiques.")
    else:
        messages.success(request, "Offre publique mise à jour manuellement.")
    return redirect("job_agent:review_detected_offers")


@login_required
def interview_coach(request, lead_id: int):
    lead = get_object_or_404(JobLead, id=lead_id, user=request.user)
    if not has_candidate_access(request.user):
        return _premium_agent_required(request)

    profile, _ = CandidateProfile.objects.get_or_create(user=request.user)
    docs, _ = CandidateDocuments.objects.get_or_create(user=request.user)
    pack = getattr(lead, "pack", None)
    if pack is None:
        pack = build_or_update_pack(user=request.user, lead=lead, profile=profile, docs=docs)

    guide = (pack.suggested_answers or {}).get("Guide entretien complet", "")
    quick_answers = {
        key: value
        for key, value in (pack.suggested_answers or {}).items()
        if key != "Guide entretien complet"
    }

    return render(
        request,
        "job_agent/interview_coach.html",
        {
            "lead": lead,
            "pack": pack,
            "guide": guide,
            "quick_answers": quick_answers,
            "menu_pack_lead_id": lead.id,
        },
    )


@login_required
@transaction.atomic
def import_public_offer(request, offer_id: int):
    offer = get_object_or_404(PublicJobOffer, id=offer_id, is_active=True)
    description = offer.description_text or ""
    if offer.skills_keywords:
        description = f"{description}\n\nMots-clés: {offer.skills_keywords}".strip()

    try:
        lead = JobLead.objects.create(
            user=request.user,
            search=_latest_search(request.user),
            url=offer.url,
            source=offer.source,
            title=offer.title,
            company=offer.company,
            location=offer.location,
            description_text=description,
            status=JobLead.STATUS_FOUND,
        )
    except IntegrityError:
        messages.info(request, "Tu as déjà importé cette offre.")
        return redirect("job_agent:lead_list")

    messages.success(request, "Offre importée dans ton espace.")
    return redirect("job_agent:lead_detail", lead_id=lead.id)


@login_required
@transaction.atomic
def adapt_public_offer(request, offer_id: int):
    offer = get_object_or_404(PublicJobOffer, id=offer_id, is_active=True)
    if not has_candidate_access(request.user):
        return _premium_agent_required(request)
    profile, _ = CandidateProfile.objects.get_or_create(user=request.user)
    docs, _ = CandidateDocuments.objects.get_or_create(user=request.user)
    description = offer.description_text or ""
    if offer.skills_keywords:
        description = f"{description}\n\nMots-clés: {offer.skills_keywords}".strip()

    lead, created = JobLead.objects.get_or_create(
        user=request.user,
        url=offer.url,
        defaults={
            "search": _latest_search(request.user),
            "source": offer.source,
            "title": offer.title,
            "company": offer.company,
            "location": offer.location,
            "description_text": description,
            "status": JobLead.STATUS_TO_APPLY,
        },
    )

    if not created:
        updated_fields = []
        for field, value in {
            "source": offer.source,
            "title": offer.title,
            "company": offer.company,
            "location": offer.location,
            "description_text": description,
        }.items():
            if value and not getattr(lead, field):
                setattr(lead, field, value)
                updated_fields.append(field)
        if lead.status == JobLead.STATUS_FOUND:
            lead.status = JobLead.STATUS_TO_APPLY
            updated_fields.append("status")
        if updated_fields:
            lead.save(update_fields=updated_fields)

    build_or_update_pack(user=request.user, lead=lead, profile=profile, docs=docs)
    if not (docs.cv_text or "").strip():
        messages.warning(
            request,
            "Ajoutez le texte de votre CV dans Documents CV pour obtenir une adaptation plus précise.",
        )
    messages.success(request, "Offre prête. Vérifiez le pack avant de postuler.")
    return redirect("job_agent:apply_wizard", lead_id=lead.id)


# ======================================================
# Ajout en masse (ultra rapide)
# ======================================================
@login_required
@transaction.atomic
def lead_bulk_add(request):
    """
    L'utilisateur colle plusieurs offres.
    - Par défaut: blocs séparés par '---'
    - Si pas de '---': accepte plusieurs URL (une par ligne) comme import rapide
    """
    if request.method == "POST":
        form = JobLeadBulkAddForm(request.POST)
        if form.is_valid():
            raw = (form.cleaned_data.get("payload") or "").strip()
            default_source = (form.cleaned_data.get("default_source") or "").strip()

            if not raw:
                messages.error(request, "Colle au moins une offre.")
                return redirect("job_agent:lead_bulk_add")

            if "---" in raw:
                blocks = [b.strip() for b in raw.split("---") if b.strip()]
            else:
                # mode simple: une URL par ligne
                urls = [l.strip() for l in raw.splitlines() if l.strip()]
                blocks = [f"URL: {u}" for u in urls]

            last_search = _latest_search(request.user)
            docs, _ = CandidateDocuments.objects.get_or_create(user=request.user)

            created = 0
            skipped = 0

            for b in blocks:
                url = _safe_get_line(b, "url")
                if not url:
                    skipped += 1
                    continue

                title = _safe_get_line(b, "titre")
                company = _safe_get_line(b, "entreprise")
                location = _safe_get_line(b, "lieu")
                source = _safe_get_line(b, "source") or default_source
                desc = _extract_description(b)

                try:
                    lead = JobLead.objects.create(
                        user=request.user,
                        search=last_search,
                        url=url,
                        source=source,
                        title=title,
                        company=company,
                        location=location,
                        description_text=desc,
                        status=JobLead.STATUS_FOUND,
                    )
                except IntegrityError:
                    skipped += 1
                    continue

                # scoring auto si description + CV texte
                if (lead.description_text or "").strip():
                    keywords = last_search.keywords if last_search else ""
                    res = heuristic_match(docs.cv_text or "", lead.description_text or "", keywords=keywords)
                    lead.match_score = res.score
                    lead.match_summary = res.summary
                    if lead.match_score >= 60:
                        lead.status = JobLead.STATUS_TO_APPLY
                    lead.save(update_fields=["match_score", "match_summary", "status"])

                created += 1

            if created:
                messages.success(request, f"{created} offres importées.")
            if skipped:
                messages.info(request, f"{skipped} offres ignorées (URL manquante ou doublons).")

            return redirect("job_agent:lead_list")
    else:
        form = JobLeadBulkAddForm()

    return render(
        request,
        "job_agent/lead_bulk_add.html",
        {
            "form": form,
            "menu_pack_lead_id": _menu_pack_lead_id_for_user(request.user),
        },
    )


# ======================================================
# Kanban (suivi rapide)
# ======================================================
@login_required
def kanban(request):
    leads = JobLead.objects.filter(user=request.user).order_by("-updated_at", "-created_at")

    grouped = {
        JobLead.STATUS_FOUND: [],
        JobLead.STATUS_TO_APPLY: [],
        JobLead.STATUS_APPLIED: [],
        JobLead.STATUS_FOLLOWUP: [],
        JobLead.STATUS_REPLY: [],
    }
    for l in leads:
        grouped.setdefault(l.status, []).append(l)

    return render(
        request,
        "job_agent/kanban.html",
        {
            "grouped": grouped,
            "menu_pack_lead_id": _menu_pack_lead_id_for_user(request.user),
            "leads_url": reverse("job_agent:lead_list"),
            "dashboard_url": reverse("job_agent:dashboard"),
        },
    )


@login_required
@transaction.atomic
def kanban_move(request, lead_id: int):
    lead = get_object_or_404(JobLead, id=lead_id, user=request.user)

    if request.method == "POST":
        new_status = (request.POST.get("status") or "").strip()
        valid = {c[0] for c in JobLead.STATUS_CHOICES}
        if new_status in valid:
            lead.status = new_status

            # ✅ si déplacé en "Postulée" => date candidature
            if new_status == JobLead.STATUS_APPLIED and not lead.applied_at:
                lead.applied_at = timezone.now()

            lead.save(update_fields=["status", "applied_at"])
        else:
            messages.error(request, "Statut invalide.")

    return redirect("job_agent:kanban")

@login_required
@transaction.atomic
def apply_wizard(request, lead_id: int):
    lead = get_object_or_404(JobLead, id=lead_id, user=request.user)
    if not has_candidate_access(request.user):
        return _premium_agent_required(request)
    profile, _ = CandidateProfile.objects.get_or_create(user=request.user)
    docs, _ = CandidateDocuments.objects.get_or_create(user=request.user)

    # 1) Pack existant ou génération si incomplet
    pack, _ = ApplicationPack.objects.get_or_create(user=request.user, lead=lead)

    pack_missing = (
        not (pack.generated_letter or "").strip()
        or not (pack.generated_email or "").strip()
        or not (pack.email_subject or "").strip()
    )
    if pack_missing:
        pack = build_or_update_pack(user=request.user, lead=lead, profile=profile, docs=docs)

    # 2) Helpers safe pour éviter crash si champs non présents
    def safe_attr(obj, name: str, default: str = ""):
        return getattr(obj, name, default) or default

    # 3) Draft (pré-remplissage)
    draft = {
        "full_name": (profile.full_name or request.user.get_username()),
        "email": (safe_attr(profile, "email") or request.user.email),
        "phone": safe_attr(profile, "phone"),
        "linkedin": safe_attr(profile, "linkedin_url"),
        "portfolio": (safe_attr(docs, "portfolio_url") or safe_attr(profile, "portfolio_url")),
        "cv_text": (docs.cv_text or ""),
        "tailored_cv_text": (pack.tailored_cv_text or ""),
        "ats_score": pack.ats_score or 0,
        "matched_keywords": pack.matched_keywords or [],
        "missing_keywords": pack.missing_keywords or [],
        "coach_notes": pack.coach_notes or "",
        "ai_status": pack.ai_status or "",
        "email_subject": (pack.email_subject or ""),
        "email_body": (pack.generated_email or ""),
        "letter": (pack.generated_letter or ""),
        "answers": (pack.suggested_answers or {}),
    }

    # 4) Sauvegarde des corrections utilisateur
    if request.method == "POST" and (request.POST.get("action") == "save_pack"):
        pack.tailored_cv_text = (request.POST.get("tailored_cv_text") or "").strip()
        pack.generated_letter = (request.POST.get("generated_letter") or "").strip()
        pack.email_subject = (request.POST.get("email_subject") or "").strip()
        pack.generated_email = (request.POST.get("generated_email") or "").strip()
        pack.coach_notes = (request.POST.get("coach_notes") or pack.coach_notes or "").strip()
        pack.save(
            update_fields=[
                "tailored_cv_text",
                "generated_letter",
                "email_subject",
                "generated_email",
                "coach_notes",
                "updated_at",
            ]
        )
        messages.success(request, "Corrections enregistrées ✅")
        return redirect("job_agent:apply_wizard", lead_id=lead.id)

    # 5) Validation : marquer comme postulée
    if request.method == "POST" and (request.POST.get("action") == "mark_applied"):
        lead.status = JobLead.STATUS_APPLIED
        if not lead.applied_at:
            lead.applied_at = timezone.now()
        lead.save(update_fields=["status", "applied_at"])
        messages.success(request, "Candidature marquée comme postulée ✅")
        return redirect("job_agent:lead_detail", lead_id=lead.id)

    return render(
        request,
        "job_agent/apply_wizard.html",
        {
            "lead": lead,
            "profile": profile,
            "docs": docs,
            "pack": pack,
            "draft": draft,
            "menu_pack_lead_id": lead.id,
        },
    )


def _pack_pdf_response(*, pack: ApplicationPack, kind: str) -> HttpResponse:
    labels = {
        "cv": ("CV adapté", pack.tailored_cv_text, "cv_adapte"),
        "letter": ("Lettre de motivation", pack.generated_letter, "lettre_motivation"),
        "email": (
            "Email de candidature",
            f"Objet: {pack.email_subject}\n\n{pack.generated_email}".strip(),
            "email_candidature",
        ),
    }
    if kind not in labels:
        raise Http404("Document introuvable.")

    title, body, slug = labels[kind]
    body = (body or "").strip() or "Document vide. Revenez au pack pour générer ou compléter ce contenu."

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        from xml.sax.saxutils import escape

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=1.7 * cm,
            leftMargin=1.7 * cm,
            topMargin=1.5 * cm,
            bottomMargin=1.5 * cm,
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "Imm97Title",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            textColor="#0F2040",
            fontSize=20,
            leading=24,
            spaceAfter=14,
        )
        meta_style = ParagraphStyle(
            "Imm97Meta",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            textColor="#D4A843",
            fontSize=10,
            leading=13,
            spaceAfter=10,
        )
        body_style = ParagraphStyle(
            "Imm97Body",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=15,
            spaceAfter=7,
        )

        story = [
            Paragraph("Immigration97 - Pack candidature", meta_style),
            Paragraph(escape(title), title_style),
        ]
        lead = pack.lead
        meta = f"{lead.title or 'Offre'} - {lead.company or 'Entreprise'}"
        story.append(Paragraph(escape(meta), meta_style))
        if kind == "cv":
            story.append(Paragraph(f"Score ATS: {pack.ats_score or 0}/100", meta_style))
        story.append(Spacer(1, 10))
        for line in body.splitlines():
            story.append(Paragraph(escape(line) if line.strip() else "&nbsp;", body_style))
        doc.build(story)
        pdf = buffer.getvalue()
    except Exception:
        raise Http404("Export PDF indisponible.")

    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{slug}_{pack.lead_id}.pdf"'
    return response


def _pack_docx_response(*, pack: ApplicationPack, kind: str) -> HttpResponse:
    labels = {
        "cv": ("CV adapté", pack.tailored_cv_text, "cv_adapte"),
        "letter": ("Lettre de motivation", pack.generated_letter, "lettre_motivation"),
        "email": (
            "Email de candidature",
            f"Objet: {pack.email_subject}\n\n{pack.generated_email}".strip(),
            "email_candidature",
        ),
    }
    if kind not in labels:
        raise Http404("Document introuvable.")

    title, body, slug = labels[kind]
    body = (body or "").strip() or "Document vide. Revenez au pack pour générer ou compléter ce contenu."

    try:
        from docx import Document
        from docx.shared import Pt

        document = Document()
        document.core_properties.title = title
        document.core_properties.subject = "Immigration97 - Pack candidature"
        document.add_paragraph("Immigration97 - Pack candidature")
        document.add_heading(title, level=1)
        lead = pack.lead
        document.add_paragraph(f"{lead.title or 'Offre'} - {lead.company or 'Entreprise'}")
        if kind == "cv":
            document.add_paragraph(f"Score ATS: {pack.ats_score or 0}/100")
        document.add_paragraph("")

        for line in body.splitlines():
            document.add_paragraph(line if line.strip() else "")

        for style_name in ["Normal"]:
            style = document.styles[style_name]
            style.font.name = "Arial"
            style.font.size = Pt(10.5)

        buffer = BytesIO()
        document.save(buffer)
        content = buffer.getvalue()
    except Exception:
        raise Http404("Export Word indisponible.")

    response = HttpResponse(
        content,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{slug}_{pack.lead_id}.docx"'
    return response


@login_required
def pack_download_pdf(request, lead_id: int, kind: str):
    lead = get_object_or_404(JobLead, id=lead_id, user=request.user)
    pack = get_object_or_404(ApplicationPack, lead=lead, user=request.user)
    return _pack_pdf_response(pack=pack, kind=kind)


@login_required
def pack_download_docx(request, lead_id: int, kind: str):
    lead = get_object_or_404(JobLead, id=lead_id, user=request.user)
    pack = get_object_or_404(ApplicationPack, lead=lead, user=request.user)
    return _pack_docx_response(pack=pack, kind=kind)


from django.http import JsonResponse
from django.views.decorators.http import require_GET

from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_GET

@login_required
@require_GET
def indeed_autofill_api(request, lead_id: int):
    # Récupération sécurisée des objets
    lead = get_object_or_404(JobLead, id=lead_id, user=request.user)
    profile, _ = CandidateProfile.objects.get_or_create(user=request.user)
    docs, _ = CandidateDocuments.objects.get_or_create(user=request.user)

    # Logique de pack d'application
    pack, _ = ApplicationPack.objects.get_or_create(user=request.user, lead=lead)
    if not pack.generated_letter or not pack.generated_email:
        pack = build_or_update_pack(user=request.user, lead=lead, profile=profile, docs=docs)

    # Construction de l'URL du CV
    cv_url = ""
    try:
        if docs.cv_file:
            cv_url = request.build_absolute_uri(docs.cv_file.url)
    except Exception:
        pass

    # Structuration des données
    data = {
        "candidate": {
            "full_name": profile.full_name or request.user.get_username(),
            "email": request.user.email or getattr(profile, "email", ""),
            "phone": getattr(request.user, "phone", "") or getattr(profile, "phone", ""),
            "city": getattr(profile, "city", ""),
            "linkedin": getattr(profile, "linkedin_url", ""),
            "portfolio": getattr(docs, "portfolio_url", "") or getattr(profile, "portfolio_url", ""),
            "cv_url": cv_url,
        },
        "application": {
            "cover_letter": pack.generated_letter or "",
            "email_subject": pack.email_subject or "",
            "answers": pack.suggested_answers or {},
        }
    }

    return JsonResponse(data)
